"""
Modul pre spracovanie dokumentov - PDF, DOCX, TXT, MD súborov.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Any


class DocumentProcessor:
    """Trieda pre spracovanie rôznych typov dokumentov."""
    
    def __init__(self):
        """Inicializácia procesora dokumentov."""
        self.logger = logging.getLogger(__name__)
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Spracuje dokument a vráti jeho obsah a metadáta.
        
        Args:
            file_path: Cesta k súboru
            
        Returns:
            Slovník s obsahom a metadátami
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": f"Súbor {file_path} neexistuje"}
        
        if file_path.suffix.lower() not in self.supported_formats:
            return {"error": f"Nepodporovaný formát súboru: {file_path.suffix}"}
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._process_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._process_docx(file_path)
            elif file_path.suffix.lower() in ['.txt', '.md']:
                return self._process_text(file_path)
        except Exception as e:
            self.logger.error(f"Chyba pri spracovaní súboru {file_path}: {e}")
            return {"error": f"Chyba pri spracovaní: {str(e)}"}
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Spracuje PDF súbor."""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    "content": text.strip(),
                    "page_count": len(reader.pages),
                    "format": "pdf",
                    "file_name": file_path.name
                }
        except ImportError:
            return {"error": "PyPDF2 nie je nainštalovaný. Použite: pip install PyPDF2"}
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Spracuje DOCX súbor."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            return {
                "content": text.strip(),
                "paragraph_count": len(doc.paragraphs),
                "format": "docx",
                "file_name": file_path.name
            }
        except ImportError:
            return {"error": "python-docx nie je nainštalovaný. Použite: pip install python-docx"}
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Spracuje textové súbory (TXT, MD)."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
            return {
                "content": content,
                "character_count": len(content),
                "line_count": len(content.split('\n')),
                "format": file_path.suffix[1:],
                "file_name": file_path.name
            }
    
    def handle_document_query(self, command: str) -> str:
        """
        Spracuje príkaz týkajúci sa dokumentov.
        
        Args:
            command: Príkaz od používateľa
            
        Returns:
            Odpoveď pre používateľa
        """
        if "analyze" in command.lower():
            return "Pre analýzu dokumentu zadajte cestu k súboru. Podporované formáty: PDF, DOCX, TXT, MD"
        elif "summarize" in command.lower():
            return "Pre zhrnutie dokumentu najprv nahrajte súbor pomocou príkazu 'analyze [cesta_k_súboru]'"
        else:
            return "Dostupné príkazy: analyze [súbor], summarize [súbor], process [súbor]"